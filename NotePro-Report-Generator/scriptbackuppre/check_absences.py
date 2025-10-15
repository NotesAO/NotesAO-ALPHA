import pandas as pd
from docx import Document
import os
from datetime import datetime
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
import subprocess
import argparse

def find_clients_with_consecutive_absences_and_no_recent_attendance(df, consecutive_count):
    clients_to_exit = []

    for index, row in df.iterrows():
        p_dates = [row[f'P{i}'] for i in range(1, 36) if f'P{i}' in df.columns and not pd.isna(row[f'P{i}'])]
        a_dates = [row[f'A{i}'] for i in range(1, 29) if f'A{i}' in df.columns and not pd.isna(row[f'A{i}'])]

        # Convert date strings to datetime objects
        p_dates = [pd.to_datetime(date) for date in p_dates]
        a_dates = [pd.to_datetime(date) for date in a_dates]

        all_dates = [(date, 'P') for date in p_dates] + [(date, 'A') for date in a_dates]
        all_dates.sort()

        count = 0
        recent_absences = []
        for date, label in reversed(all_dates):
            if label == 'A':
                count += 1
                recent_absences.append(date)
                if count == consecutive_count:
                    break
            else:
                count = 0
                recent_absences = []

        if count == consecutive_count:
            last_absence_date = recent_absences[0]
            subsequent_attendance = [date for date, label in all_dates if label == 'P' and date > last_absence_date]
            if not subsequent_attendance:
                clients_to_exit.append(row['client_id'])

    return clients_to_exit

def get_client_names_by_program(df, client_ids):
    program_groups = df[df['client_id'].isin(client_ids)].groupby('program_name')
    program_clients = {}
    for program, group in program_groups:
        names = group[['first_name', 'last_name']].apply(lambda x: f"{x['first_name']} {x['last_name']}", axis=1).tolist()
        names.sort(key=lambda name: name.split()[-1])  # Sort by last name
        program_clients[program] = names
    return program_clients

def find_clients_with_no_attendance_since_intake(df, report_date):
    report_date = pd.to_datetime(report_date, format='%m/%d/%Y')
    recent_intakes = df[
        (pd.to_datetime(df['orientation_date'], errors='coerce') <= report_date - pd.Timedelta(days=8)) &
        (df['last_attended'].isna())
    ]
    client_ids = recent_intakes['client_id'].tolist()
    return client_ids

def save_results_to_doc(results, output_path):
    doc = Document()
    doc.add_heading('Absence Report', level=1)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.orientation = WD_ORIENT.PORTRAIT

    for title, result in results.items():
        if title == "Clients with 4 Absences":
            #doc.add_page_break()
            doc.add_heading('Absence Report Continued - Totals', level=1)

        doc.add_heading(title, level=2)
        for program, names in result.items():
            doc.add_heading(program, level=3)
            if names:
                doc.add_paragraph(", ".join(names))
            else:
                doc.add_paragraph("No clients found.")

    doc.save(output_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    # Use 'soffice' for LibreOffice conversion
    subprocess.run([
        'libreoffice',
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', os.path.dirname(pdf_path),
        docx_path
    ], check=True)

def process_absences(csv_file, report_date, output_directory):
    df = pd.read_csv(csv_file)
    df = df[df['exit_reason'] == 'Not Exited']

    # Ensure necessary columns are present
    required_columns = ['client_id', 'first_name', 'last_name', 'program_name', 'orientation_date', 'last_attended', 'absence_unexcused']
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Missing required columns in CSV file: {', '.join(missing_columns)}")

    clients_with_3_consecutive_absences = find_clients_with_consecutive_absences_and_no_recent_attendance(df, 3)
    program_clients_3_consecutive = get_client_names_by_program(df, clients_with_3_consecutive_absences)

    clients_with_2_consecutive_absences = find_clients_with_consecutive_absences_and_no_recent_attendance(df, 2)
    clients_with_2_consecutive_absences_unique = [
        client for client in clients_with_2_consecutive_absences if client not in clients_with_3_consecutive_absences
    ]
    program_clients_2_consecutive = get_client_names_by_program(df, clients_with_2_consecutive_absences_unique)

    clients_with_no_attendance_since_intake = find_clients_with_no_attendance_since_intake(df, report_date)
    program_clients_no_attendance_since_intake = get_client_names_by_program(df, clients_with_no_attendance_since_intake)

    clients_with_4_unexcused_absences = df[df['absence_unexcused'] == 4]['client_id'].tolist()
    program_clients_4_absences = get_client_names_by_program(df, clients_with_4_unexcused_absences)

    clients_with_5_or_more_unexcused_absences = df[df['absence_unexcused'] >= 5]['client_id'].tolist()
    program_clients_5_or_more_absences = get_client_names_by_program(df, clients_with_5_or_more_unexcused_absences)

    results = {
        "Clients with 3 Consecutive Absences (Exit)": program_clients_3_consecutive,
        "Clients with 2 Consecutive Absences (Warning)": program_clients_2_consecutive,
        "Clients who have NOT Attended Since Intake/Orientation": program_clients_no_attendance_since_intake,
        "Clients with 4 Absences": program_clients_4_absences,
        "Clients with 5 or More Absences": program_clients_5_or_more_absences,
    }

    today = datetime.now().strftime('%Y%m%d')
    output_path_docx = os.path.join(output_directory, f'AbsenceReport_{today}.docx')

    # Ensure the directory exists
    os.makedirs(os.path.dirname(output_path_docx), exist_ok=True)

    save_results_to_doc(results, output_path_docx)

    output_path_pdf = os.path.join(output_directory, f'AbsenceReport_{today}.pdf')
    convert_docx_to_pdf(output_path_docx, output_path_pdf)
    os.remove(output_path_docx)

    return output_path_pdf

def main():
    parser = argparse.ArgumentParser(description='Process absences and generate reports.')
    parser.add_argument('--csv_file', required=True, help='The input CSV file path')
    parser.add_argument('--report_date', required=False, help='The report date', default=datetime.now().strftime('%m/%d/%Y'))
    parser.add_argument('--output_directory', required=True, help='The output directory for the report')

    args = parser.parse_args()

    csv_file = args.csv_file
    report_date = args.report_date
    output_directory = args.output_directory

    # Ensure the output directory exists
    os.makedirs(output_directory, exist_ok=True)

    try:
        output_pdf = process_absences(csv_file, report_date, output_directory)
        print(f"Absence report generated: {output_pdf}")
    except Exception as e:
        print(f"Error generating absence report: {e}")

if __name__ == '__main__':
    main()
